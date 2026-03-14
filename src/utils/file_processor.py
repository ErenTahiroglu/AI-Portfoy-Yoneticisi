"""
🧩 Puzzle Parça: Dosya İşleme (Excel/DOCX/PDF)
=================================================
Kullanıcının yüklediği dosyalardan ticker çıkarır,
analiz sonuçlarını Excel/PDF/DOCX formatında dışa aktarır.
"""

import logging
import os
import re
import io

import pandas as pd
import docx
from fpdf import FPDF

logger = logging.getLogger(__name__)


# ══════════════════════════════════════════════════════════════════════════════
# Ticker Çıkarma
# ══════════════════════════════════════════════════════════════════════════════

def extract_tickers_from_text(text):
    """Metinden hisse/fon sembollerini çıkarır (AAPL, TP2, THYAO.IS vb.)."""
    words = re.findall(r'\b[A-Za-z][A-Za-z0-9.]*\b', text)
    return list(set([word.upper() for word in words]))


def process_uploaded_file(uploaded_file):
    """Yüklenen dosyadan ticker listesi çıkarır."""
    filename = uploaded_file.name.lower()
    tickers = []
    try:
        if filename.endswith('.txt'):
            tickers = extract_tickers_from_text(uploaded_file.read().decode('utf-8'))
        elif filename.endswith('.docx'):
            doc = docx.Document(uploaded_file)
            tickers = extract_tickers_from_text(" ".join([p.text for p in doc.paragraphs]))
        elif filename.endswith('.csv'):
            df = pd.read_csv(uploaded_file, sep=None, engine='python', on_bad_lines='skip')
            tickers = extract_tickers_from_text(df.to_string())
        elif filename.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(uploaded_file)
            tickers = extract_tickers_from_text(df.to_string())
    except Exception as e:
        logger.error(f"Dosya okunurken hata: {e}")
        raise ValueError(f"Dosya okuma hatası: {e}")
    return tickers


# ══════════════════════════════════════════════════════════════════════════════
# Excel Dışa Aktarım
# ══════════════════════════════════════════════════════════════════════════════

def to_excel(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Arindirma_Analizi')
    return output.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# PDF Dışa Aktarım — fpdf2 Unicode (Türkçe tam destek)
# ══════════════════════════════════════════════════════════════════════════════

def _get_unicode_font_path():
    """Sistemde bulunan Unicode destekli bir font yolu döndürür."""
    # Windows
    win_fonts = [
        os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts", f)
        for f in ["arial.ttf", "segoeui.ttf", "tahoma.ttf", "calibri.ttf"]
    ]
    # macOS / Linux
    unix_fonts = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    ]
    
    for path in win_fonts + unix_fonts:
        if os.path.exists(path):
            return path
    return None


def _clean_emojis(text: str) -> str:
    """PDF fontlarının desteklemediği emojileri anlamlı metne çevirir."""
    emoji_map = {
        '✅': '[OK]', '⚠️': '[!]', '❌': '[X]', '🔍': '[?]',
        '📈': '[^]', '📊': '[#]', '💰': '[$]', '🎯': '[*]',
        '🤖': '[AI]', '☪️': '[I]', '🌍': '[W]', '🛡️': '[S]',
        '📂': '[F]', '⚡': '[!]', '🌐': '[W]', '🚀': '[>]',
        '🧩': '[P]', '🔒': '[L]', '♻️': '[R]',
    }
    for emoji, replacement in emoji_map.items():
        text = text.replace(emoji, replacement)
    # Kalan unicode emojileri temizle (BMP dışı karakterler)
    return ''.join(c for c in text if ord(c) < 0x10000 or c in 'ğüşöçıİĞÜŞÖÇ₺')


def create_pdf(report_text):
    """PDF oluşturur — fpdf2 Unicode font ile Türkçe tam destek."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Unicode font yükle (Türkçe harfler korunacak)
    font_path = _get_unicode_font_path()
    if font_path:
        pdf.add_font("UnicodeFont", "", font_path)
        pdf.add_font("UnicodeFont", "B", font_path)
        font_name = "UnicodeFont"
    else:
        # Fallback: Arial (Latin-1 only)
        font_name = "Arial"
        logger.warning("Unicode font bulunamadı, ASCII fallback kullanılıyor.")
    
    # Emojileri temizle (fontlar desteklemiyor)
    clean_text = _clean_emojis(report_text)
    
    for line in clean_text.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(4)
            continue
            
        # Markdown bold sembollerini temizle
        clean_line = line.replace('**', '')
        
        # ANA BAŞLIKLAR (#)
        if clean_line.startswith('# '):
            pdf.set_font(font_name, 'B', 16)
            pdf.set_text_color(0, 51, 102)
            pdf.multi_cell(0, 10, clean_line.replace('# ', ''))
            pdf.ln(2)
            
        # ALT BAŞLIKLAR (##)
        elif clean_line.startswith('## '):
            pdf.set_font(font_name, 'B', 13)
            pdf.set_text_color(0, 0, 0)
            pdf.multi_cell(0, 8, clean_line.replace('## ', ''))
            pdf.ln(1)
            
        # MADDE İŞARETLERİ (-)
        elif clean_line.startswith('- '):
            pdf.set_font(font_name, '', 11)
            pdf.set_text_color(50, 50, 50)
            pdf.set_x(18)  # Set a defined left margin for bullets instead of using spaces
            pdf.multi_cell(0, 6, clean_line)
            pdf.set_x(15)  # Reset to default
            
        # NORMAL PARAGRAF METNİ
        else:
            pdf.set_font(font_name, '', 11)
            pdf.set_text_color(0, 0, 0)
            pdf.set_x(15)
            pdf.multi_cell(0, 6, clean_line)
    
    out = pdf.output(dest="S")
    if isinstance(out, str):
        return out.encode("latin-1")
    return bytes(out)


# ══════════════════════════════════════════════════════════════════════════════
# DOCX Dışa Aktarım
# ══════════════════════════════════════════════════════════════════════════════

def create_docx(report_text):
    """Word formatı — emojiler ve Türkçe harfler tam korunur."""
    doc = docx.Document()
    doc.add_heading('Portföy Analiz Raporu', 0)
    for line in report_text.split('\n'):
        if line.startswith('## '):
            doc.add_heading(line.replace('##', '').strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('#', '').strip(), level=1)
        elif line.startswith('- '):
            doc.add_paragraph(line.replace('-', '').strip(), style='List Bullet')
        elif line.strip() != "":
            clean_line = line.replace('**', '')
            doc.add_paragraph(clean_line)
            
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()