import pandas as pd
import docx
import re
import io

from fpdf import FPDF

def extract_tickers_from_text(text):
    words = re.findall(r'\b[A-Za-z]+\b', text)
    return list(set([word.upper() for word in words]))

def process_uploaded_file(uploaded_file):
    filename = uploaded_file.name.lower()
    tickers = []
    try:
        if filename.endswith('.txt'):
            tickers = extract_tickers_from_text(uploaded_file.read().decode('utf-8'))
        elif filename.endswith('.docx'):
            doc = docx.Document(uploaded_file)
            tickers = extract_tickers_from_text(" ".join([p.text for p in doc.paragraphs]))
        elif filename.endswith('.csv'):
            df = pd.read_csv(uploaded_file, sep=None, engine='python', on_bad_lines='skip')
            tickers = extract_tickers_from_text(df.to_string())
        elif filename.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(uploaded_file)
            tickers = extract_tickers_from_text(df.to_string())
    except Exception as e:
        print(f"Dosya okunurken hata: {e}")
        raise ValueError(f"Dosya okuma hatası: {e}")
    return tickers

def to_excel(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Arindirma_Analizi')
    return output.getvalue()

def create_pdf(report_text):
    """PDF Formatı emojileri desteklemez, Markdown'ı okunaklı başlık ve paragraflara çevirir"""
    replacements = {
        'ı': 'i', 'İ': 'I', 'ş': 's', 'Ş': 'S', 'ğ': 'g', 'Ğ': 'G',
        'ç': 'c', 'Ç': 'C', 'ö': 'o', 'Ö': 'O', 'ü': 'u', 'Ü': 'U', '₺': 'TL'
    }
    clean_text = report_text
    for k, v in replacements.items():
        clean_text = clean_text.replace(k, v)
    
    # Emojileri temizle
    clean_text = clean_text.encode('ascii', 'ignore').decode('ascii')
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    for line in clean_text.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(4) # Paragraflar arasına nefes alma boşluğu
            continue
            
        # Göze batan ** sembollerini temizliyoruz
        clean_line = line.replace('**', '')
        
        # ANA BAŞLIKLAR (#)
        if clean_line.startswith('# '):
            pdf.set_font("Arial", 'B', 16)
            pdf.set_text_color(0, 51, 102) # Kurumsal Lacivert renk
            pdf.multi_cell(0, 10, clean_line.replace('# ', ''))
            pdf.ln(2)
            
        # ALT BAŞLIKLAR (##)
        elif clean_line.startswith('## '):
            pdf.set_font("Arial", 'B', 13)
            pdf.set_text_color(0, 0, 0) # Siyah
            pdf.multi_cell(0, 8, clean_line.replace('## ', ''))
            pdf.ln(1)
            
        # MADDE İŞARETLERİ (-)
        elif clean_line.startswith('- '):
            pdf.set_font("Arial", '', 11)
            pdf.set_text_color(50, 50, 50) # Koyu Gri
            pdf.multi_cell(0, 6, "  " + clean_line)
            
        # NORMAL PARAGRAF METNİ
        else:
            pdf.set_font("Arial", '', 11)
            pdf.set_text_color(0, 0, 0)
            pdf.multi_cell(0, 6, clean_line)
    
    out = pdf.output(dest="S")
    if isinstance(out, str):
        return out.encode("latin-1")
    return bytes(out)

def create_docx(report_text):
    """Word formatı emojileri ve Türkçe harfleri mükemmel şekilde korur"""
    doc = docx.Document()
    doc.add_heading('AI İslami Portföy Yöneticisi Raporu', 0)
    for line in report_text.split('\n'):
        if line.startswith('## '):
            doc.add_heading(line.replace('##', '').strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('#', '').strip(), level=1)
        elif line.startswith('- '):
            doc.add_paragraph(line.replace('-', '').strip(), style='List Bullet')
        elif line.strip() != "":
            # Word'de de Markdown sembollerini temizleyelim
            clean_line = line.replace('**', '')
            doc.add_paragraph(clean_line)
            
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()